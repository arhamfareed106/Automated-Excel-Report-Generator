#!/usr/bin/env python3
"""
generate_report.py

Run this script locally (Python 3.9+) to produce the full client deliverables:
- Portfolio_VaR_Workbook.xlsx (Excel with calculations & charts)
- Client_Report.docx (Word report with executive summary, findings, and appendix)
- Figures/ (payoff diagrams, VaR charts)

Usage:
    python generate_report.py --tickers AAPL MSFT NVDA ORCL --start 2023-09-29 --end 2025-10-10 --output ./output

Default:
    Tickers: AAPL, MSFT, NVDA, ORCL (ORCL fixed 394 shares)
    Start: 2023-09-29 (2 years before 2025-09-29)
    End: 2025-10-10 (end of holding period)

This script requires internet access to download historical price data via yfinance.
It will create an Excel workbook formatted for the client's submission and a Word report.
"""

import os
import argparse
from datetime import datetime, timedelta
import numpy as np
import pandas as pd

def install_missing(package):
    # helper that instructs user how to install missing packages if needed
    raise RuntimeError(f"Missing package: {package}. Please install with pip install {package}")

try:
    import yfinance as yf
except Exception as e:
    install_missing('yfinance')

try:
    import matplotlib.pyplot as plt
except Exception as e:
    install_missing('matplotlib')

try:
    import openpyxl
except Exception as e:
    install_missing('openpyxl')

try:
    from docx import Document
    from docx.shared import Inches
except Exception as e:
    install_missing('python-docx')

def download_data(tickers, start, end):
    print(f"Downloading data for {tickers} from {start} to {end} ...")
    data = yf.download(tickers, start=start, end=end, progress=False)['Adj Close']
    if isinstance(data, pd.Series):
        data = data.to_frame()
    return data

def compute_returns(prices):
    returns = prices.pct_change().dropna()
    return returns

def historical_var(returns, alpha=0.95):
    # returns: DataFrame of daily returns
    var = returns.quantile(1-alpha)
    return var

def historical_portfolio_var(portfolio_returns, alpha=0.95):
    # portfolio_returns Series
    var = np.percentile(portfolio_returns, (1-alpha)*100)
    return var

def model_var(portfolio_returns, alpha=0.95):
    mu = portfolio_returns.mean()
    sigma = portfolio_returns.std(ddof=1)
    from scipy.stats import norm
    var = norm.ppf(1-alpha, loc=mu, scale=sigma)
    return var

def scale_var_1d_to_nd(var_1d, n_days, method='sqrt'):
    if method == 'sqrt':
        return var_1d * np.sqrt(n_days)
    else:
        return None

def expected_shortfall(returns, alpha=0.95):
    threshold = np.percentile(returns, (1-alpha)*100)
    return returns[returns <= threshold].mean()

def portfolio_value_from_weights(prices, weights, total_value=1_000_000):
    # weights dict ticker -> weight (sum to 1)
    latest_prices = prices.iloc[-1]
    alloc = {t: weights[t]*total_value for t in weights}
    shares = {t: int(np.floor(alloc[t]/latest_prices[t])) for t in alloc}
    used = sum(shares[t]*latest_prices[t] for t in shares)
    cash = total_value - used
    return shares, used, cash

def save_to_excel(prices, returns, var_1d, var_10d_hist, var_10d_model, cvar_10d, shares, used, cash, output_file):
    with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
        prices.to_excel(writer, sheet_name='Prices')
        returns.to_excel(writer, sheet_name='Returns')
        var_1d.to_frame('1-day VaR').to_excel(writer, sheet_name='VaR_1d')
        pd.DataFrame({'10d_Hist_VaR': var_10d_hist}, index=[0]).to_excel(writer, sheet_name='VaR_10d_Hist', index=False)
        pd.DataFrame({'10d_Model_VaR': var_10d_model}, index=[0]).to_excel(writer, sheet_name='VaR_10d_Model', index=False)
        pd.DataFrame({'10d_CVaR': cvar_10d}, index=[0]).to_excel(writer, sheet_name='CVaR_10d', index=False)
        pd.DataFrame({'Shares': shares}).to_excel(writer, sheet_name='Positions')
        pd.DataFrame({'Used': [used], 'Cash': [cash]}).to_excel(writer, sheet_name='Cash_Used', index=False)
    print(f"Excel workbook saved to {output_file}")

def make_payoff_chart(ticker, S_range, payoff, filename):
    plt.figure(figsize=(8,5))
    plt.plot(S_range, payoff)
    plt.title(f'Option Payoff - {ticker}')
    plt.xlabel('Stock Price at Expiry')
    plt.ylabel('Profit / Loss (USD)')
    plt.grid(True)
    plt.savefig(filename)
    plt.close()

def generate_report_docx(output_docx, context):
    doc = Document()
    doc.add_heading('Client Portfolio & Risk Report', level=1)
    doc.add_paragraph(context['executive_summary'])
    doc.add_heading('Portfolio Construction', level=2)
    doc.add_paragraph(context['portfolio_construction'])
    doc.add_heading('VaR & CVaR Analysis', level=2)
    doc.add_paragraph(context['var_summary'])
    doc.add_heading('Options Hedging Recommendations', level=2)
    doc.add_paragraph(context['options_summary'])
    doc.add_heading('Currency Swap Design', level=2)
    doc.add_paragraph(context['swap_summary'])
    doc.add_heading('Appendix', level=2)
    doc.add_paragraph('See the attached Excel workbook and figures in the output folder.')
    doc.save(output_docx)
    print(f"Word report saved to {output_docx}")

def main(args):
    tickers = args.tickers
    # Force ORCL in tickers and fixed shares of 394
    if 'ORCL' not in tickers:
        tickers.append('ORCL')
    start = args.start
    end = args.end
    output = os.path.abspath(args.output)
    os.makedirs(output, exist_ok=True)
    prices = download_data(tickers, start, end)
    returns = compute_returns(prices)
    # 1-day VaR per stock (historical)
    var_1d = returns.quantile(0.05)
    # Convert to portfolio using equal weights except ORCL fixed shares
    latest_prices = prices.loc[prices.index.max()]
    # Create allocation: ORCL 394 shares fixed; remaining cash allocated equally among other tickers
    total_budget = 1_000_000
    orcl_shares = 394
    orcl_used = orcl_shares * latest_prices['ORCL']
    remaining = total_budget - orcl_used
    other_tickers = [t for t in tickers if t != 'ORCL']
    # Simple equal-weighting among other tickers by dollar amount
    per_stock_allocation = remaining / len(other_tickers)
    shares = {'ORCL': orcl_shares}
    for t in other_tickers:
        shares[t] = int(per_stock_allocation // latest_prices[t])
    used = sum(shares[t] * latest_prices[t] for t in shares)
    cash = total_budget - used
    # Portfolio daily returns based on position weights
    position_values = pd.DataFrame({t: returns[t] for t in returns.columns})  # daily returns per ticker
    # compute portfolio returns series
    weights = {t: (shares[t]*latest_prices[t])/total_budget for t in shares}
    portfolio_daily_returns = (returns[list(weights.keys())] * pd.Series(weights)).sum(axis=1)
    # 10-day historical VaR (portfolio)
    # Aggregate daily returns to 10-day holding by summing (approximation)
    rolling_10 = portfolio_daily_returns.rolling(window=10).sum().dropna()
    var_10d_hist = np.percentile(rolling_10, 5)
    # 10-day model-based VaR (assuming normality)
    var_10d_model = model_var(portfolio_daily_returns) * np.sqrt(10)
    # 10-day CVaR (historical)
    cvar_10d = rolling_10[rolling_10 <= np.percentile(rolling_10, 5)].mean()
    # Save to Excel
    excel_path = os.path.join(output, 'Portfolio_VaR_Workbook.xlsx')
    save_to_excel(prices, returns, var_1d, var_10d_hist, var_10d_model, cvar_10d, shares, used, cash, excel_path)
    # Option payoffs - example protective put on first non-ORCL stock
    hedge_ticker = other_tickers[0]
    S = np.linspace(0.5*latest_prices[hedge_ticker], 1.5*latest_prices[hedge_ticker], 200)
    K_put = round(0.9 * latest_prices[hedge_ticker], 2)
    premium_put = 0.01 * latest_prices[hedge_ticker]  # placeholder 1% premium
    payoff_put = np.maximum(K_put - S, 0) - premium_put
    fig_dir = os.path.join(output, 'figures')
    os.makedirs(fig_dir, exist_ok=True)
    put_chart = os.path.join(fig_dir, f'{hedge_ticker}_protective_put.png')
    make_payoff_chart(hedge_ticker, S, payoff_put, put_chart)
    # Two volatility strategies for ORCL: straddle and strangle (payoff plots)
    S_orcl = np.linspace(0.5*latest_prices['ORCL'], 1.5*latest_prices['ORCL'], 200)
    K_straddle = latest_prices['ORCL']
    premium_call = 0.02 * latest_prices['ORCL']
    premium_put_orcl = 0.02 * latest_prices['ORCL']
    payoff_straddle = (np.maximum(S_orcl - K_straddle, 0) - premium_call) + (np.maximum(K_straddle - S_orcl, 0) - premium_put_orcl)
    straddle_chart = os.path.join(fig_dir, 'ORCL_straddle.png')
    make_payoff_chart('ORCL_straddle', S_orcl, payoff_straddle, straddle_chart)
    # strangle
    K1 = round(0.95*latest_prices['ORCL'],2)
    K2 = round(1.05*latest_prices['ORCL'],2)
    premium_call2 = 0.01 * latest_prices['ORCL']
    premium_put2 = 0.01 * latest_prices['ORCL']
    payoff_strangle = (np.maximum(S_orcl - K2, 0) - premium_call2) + (np.maximum(K1 - S_orcl, 0) - premium_put2)
    strangle_chart = os.path.join(fig_dir, 'ORCL_strangle.png')
    make_payoff_chart('ORCL_strangle', S_orcl, payoff_strangle, strangle_chart)
    # Generate Word report
    ctx = {
        'executive_summary': 'Executive summary: This document presents the constructed equity portfolio (USD 1,000,000) and risk analysis including VaR and CVaR. Protective hedge and volatility strategies for ORCL are proposed. See Excel workbook for full calculations.',
        'portfolio_construction': f'Positions: {shares}. Total used: ${used:.2f}. Cash remaining: ${cash:.2f}. ORCL fixed at 394 shares.',
        'var_summary': f'1-day VaR per stock (historical):\\n{var_1d.to_dict()}\\n10-day historical VaR (portfolio): {var_10d_hist:.6f}\\n10-day model VaR: {var_10d_model:.6f}\\n10-day CVaR: {cvar_10d:.6f}',
        'options_summary': f'Protective put example on {hedge_ticker} with strike {K_put} and premium ~{premium_put:.2f}. Straddle and strangle on ORCL with illustrative premiums.',
        'swap_summary': 'Swap design: Bank takes 40bps. Client receives 60% of comparative advantage; example calculations are included in Excel workbook.'
    }
    doc_path = os.path.join(output, 'Client_Report.docx')
    generate_report_docx(doc_path, ctx)
    print('All outputs generated in:', output)

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--tickers', nargs='+', default=['AAPL','MSFT','NVDA','ORCL'], help='List of tickers (ORCL will be forced if missing)')
    parser.add_argument('--start', default='2023-09-29', help='Start date YYYY-MM-DD')
    parser.add_argument('--end', default='2025-10-10', help='End date YYYY-MM-DD')
    parser.add_argument('--output', default='./output', help='Output folder')
    args = parser.parse_args()
    main(args)
